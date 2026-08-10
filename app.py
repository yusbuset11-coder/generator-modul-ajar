from io import BytesIO
import json
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
import google.generativeai as genai
import streamlit as st

st.set_page_config(
    page_title="GENERATOR: MODUL AJAR PEMBELAJARAN MENDALAM",
    page_icon="📚",
    layout="wide",
)


def check_password():
  """Mengembalikan True jika pengguna memasukkan password yang benar."""

  def password_entered():
    if st.session_state["password"] == "modulpm230371":
      st.session_state["password_correct"] = True
      del st.session_state["password"]
    else:
      st.session_state["password_correct"] = False

  if "password_correct" not in st.session_state:
    st.text_input(
        "🔑 Masukkan Password Akses Aplikasi:",
        type="password",
        on_change=password_entered,
        key="password",
    )
    return False
  elif not st.session_state["password_correct"]:
    st.text_input(
        "🔑 Masukkan Password Akses Aplikasi:",
        type="password",
        on_change=password_entered,
        key="password",
    )
    st.error("😕 Maaf, password yang Anda masukkan salah.")
    return False
  else:
    return True


if not check_password():
  st.stop()
# ===================================

# Custom CSS untuk tampilan UI yang modern dan profesional
st.markdown(
    """
    <style>
    .stApp {
        background-color: #0e1117;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    .header-card {
        background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
        padding: 16px 20px;
        border-radius: 12px;
        border: 1px solid #334155;
        margin-bottom: 20px;
        box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.3);
    }
    .header-title {
        color: #f8fafc;
        font-size: 15px;
        font-weight: 700;
        margin: 0;
        letter-spacing: 0.3px;
    }
    @keyframes blink-animation {
        0% { opacity: 1; color: #facc15; }
        50% { opacity: 0.35; color: #38bdf8; }
        100% { opacity: 1; color: #facc15; }
    }
    .header-subtitle {
        font-size: 11.5px;
        margin-top: 6px;
        margin-bottom: 0;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
        animation: blink-animation 1.6s infinite ease-in-out;
        font-weight: 600;
    }
    .stButton > button {
        width: 100%;
        border-radius: 8px;
        font-weight: 600;
        background: linear-gradient(135deg, #4f46e5 0%, #3b82f6 100%);
        color: white;
        border: none;
        padding: 0.65rem 1rem;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.35);
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #4338ca 0%, #2563eb 100%);
        box-shadow: 0 6px 18px rgba(59, 130, 246, 0.5);
        transform: translateY(-2px);
    }
    [data-testid="stSidebar"] {
        background-color: #111827;
        border-right: 1px solid #1f2937;
    }
    </style>
""",
    unsafe_allow_html=True,
)

# Tampilan Header Modern dalam Card
st.markdown(
    """
    <div class="header-card">
        <h2 class="header-title">
            <marquee behavior="scroll" direction="left" scrollamount="7" style="color: #38bdf8; text-shadow: 0 0 12px rgba(56, 189, 248, 0.5);">📚 GENERATOR: MODUL AJAR PEMBELAJARAN MENDALAM</marquee>
        </h2>
        <div class="header-subtitle">
            <b>Pengembang:</b> Yustinus Budi Setyanta - PS Cabdin Bangkalan &nbsp;|&nbsp; 
            <em>Aplikasi Otomatisasi Perancangan Pembelajaran Deep Learning</em>
        </div>
    </div>
""",
    unsafe_allow_html=True,
)

# Input Pengguna di Sidebar
with st.sidebar:
  st.header("⚙️ Parameter Pembelajaran")
  api_key = st.text_input("Masukkan Google Gemini API Key", type="password")

  jenjang_pendidikan = st.selectbox(
      "Pilih Jenjang Pendidikan",
      ["SD / MI", "SMP / MTs", "SMA / MA", "SMK / MAK"],
  )

  if jenjang_pendidikan == "SD / MI":
    default_mapel = "Tematik / Kelas"
    jp_guidance = "Panduan: 1 JP = 35 Menit"
    fase_options = [
        "Fase A / Kelas 1 SD",
        "Fase A / Kelas 2 SD",
        "Fase B / Kelas 3 SD",
        "Fase B / Kelas 4 SD",
        "Fase C / Kelas 5 SD",
        "Fase C / Kelas 6 SD",
    ]
  elif jenjang_pendidikan == "SMP / MTs":
    default_mapel = "Matematika / IPA / IPS"
    jp_guidance = "Panduan: 1 JP = 40 Menit"
    fase_options = [
        "Fase D / Kelas 7 SMP",
        "Fase D / Kelas 8 SMP",
        "Fase D / Kelas 9 SMP",
    ]
  elif jenjang_pendidikan == "SMA / MA":
    default_mapel = "Bahasa Indonesia / Matematika"
    jp_guidance = "Panduan: 1 JP = 45 Menit"
    fase_options = [
        "Fase E / Kelas X SMA",
        "Fase F / Kelas XI SMA",
        "Fase F / Kelas XII SMA",
    ]
  else:
    default_mapel = (
        "Dasar-dasar Teknik Otomotif / Produk Kreatif dan Kewirausahaan"
    )
    jp_guidance = "Panduan: 1 JP = 45 Menit"
    fase_options = [
        "Fase E / Kelas X SMK (Program Dasar Keahlian)",
        "Fase F / Kelas XI SMK (Konsentrasi Keahlian)",
        "Fase F / Kelas XII SMK (Konsentrasi Keahlian)",
    ]

  mata_pelajaran = st.text_input(
      "Mata Pelajaran / Program Kejuruan", default_mapel
  )
  fase_kelas = st.selectbox("Fase / Kelas", fase_options)

  topik = st.text_input(
      "Topik / Materi Pokok / Elemen",
      (
          "Contoh: Pemeliharaan Sistem Rem Kendaraan Ringan"
          if jenjang_pendidikan == "SMK / MAK"
          else "Contoh: Menyimak Teks Laporan Observasi Secara Kritis"
      ),
  )

  st.caption(jp_guidance)
  alokasi_waktu = st.text_input(
      "Alokasi Waktu", "2 JP (2 x 45 Menit)"
  )
  pertemuan_ke = st.text_input("Pertemuan Ke-", "1 (Pertemuan Pertama)")

  st.markdown("---")
  st.header("🏫 Identitas Satuan Pendidikan")
  nama_sekolah = st.text_input("Nama Sekolah", "SMK Miftahut Tholibin Kwanyar")
  semester = st.selectbox("Semester", ["Ganjil", "Genap"])
  tahun_pelajaran = st.text_input("Tahun Pelajaran", "2025/2026")

  st.markdown("---")
  st.header("✍️ Identitas Pengesahan Dokumen")
  nama_kota = st.text_input("Nama Kota", "Bangkalan")
  tanggal_pembuatan = st.text_input(
      "Tanggal / Bulan / Tahun", "5 Agustus 2026"
  )
  nama_penulis = st.text_input(
      "Nama Penulis Modul", "Yustinus Budi Setyanta, S.Pd., M.Pd."
  )
  nip_penulis = st.text_input("NIP Penulis", "196908302005011003")


# Fungsi pembantu warna latar belakang sel tabel Word
def set_cell_background(cell, fill_color):
  shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
  cell._tc.get_or_add_tcPr().append(shading_elm)


# Fungsi pembentuk dokumen Word lengkap berbasis Tabel Matriks sesuai sistematika baru
def generate_docx(
    data_ai,
    nama_sekolah,
    semester,
    tahun_pelajaran,
    mata_pelajaran,
    fase_kelas,
    topik,
    alokasi_waktu,
    pertemuan_ke,
    nama_penulis,
    nama_kota,
    tanggal_pembuatan,
    nip_penulis,
):
  doc = docx.Document()

  # Pengaturan margin halaman
  for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

  style = doc.styles["Normal"]
  font = style.font
  font.name = "Arial"
  font.size = Pt(10)
  font.color.rgb = RGBColor(51, 51, 51)

  # Judul Utama Dokumen
  p_title = doc.add_paragraph()
  p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p_title.paragraph_format.space_before = Pt(0)
  p_title.paragraph_format.space_after = Pt(12)
  run_title = p_title.add_run("MODUL AJAR PEMBELAJARAN MENDALAM")
  run_title.font.name = "Arial"
  run_title.font.size = Pt(13)
  run_title.font.bold = True

  def add_section_table(title_text, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[0].text = title_text
    set_cell_background(hdr_cells[0], "FFE599")
    for p in hdr_cells[0].paragraphs:
      p.alignment = WD_ALIGN_PARAGRAPH.LEFT
      p.paragraph_format.space_before = Pt(4)
      p.paragraph_format.space_after = Pt(4)
      for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(51, 51, 51)

    # Content Rows
    for idx, (label, val) in enumerate(rows_data):
      row_cells = table.rows[idx + 1].cells
      row_cells[0].text = label
      row_cells[1].text = str(val)
      row_cells[0].width = Inches(2.3)
      row_cells[1].width = Inches(4.2)

      set_cell_background(row_cells[0], "F2F5F9")

      for p in row_cells[0].paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
          run.font.size = Pt(10)
          run.font.bold = True
          run.font.color.rgb = RGBColor(51, 51, 51)

      for p in row_cells[1].paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        text_p = p.text.strip()
        if title_text == "IDENTIFIKASI DAN INFORMASI UMUM":
          for run in p.runs:
            run.font.size = Pt(10)
            run.font.bold = False
            run.font.color.rgb = RGBColor(51, 51, 51)
        else:
          is_subheader = (
              text_p.startswith("Model")
              or text_p.startswith("Metode")
              or text_p.startswith("Lingkungan")
              or text_p.startswith("Ruang")
              or text_p.startswith("Tahap")
              or text_p.startswith("Kegiatan")
              or text_p.startswith("Memahami")
              or text_p.startswith("Mengaplikasi")
              or text_p.startswith("Merefleksi")
              or text_p.startswith("Asesmen")
              or text_p.startswith("1.")
              or text_p.startswith("2.")
              or text_p.startswith("3.")
              or text_p.startswith("4.")
              or text_p.endswith(":")
              or (len(text_p) < 85 and not text_p.endswith("."))
          )
          for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(51, 51, 51)
            if is_subheader:
              run.font.bold = True
            else:
              run.font.bold = False

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

  # 1. Identifikasi dan Informasi Umum
  tabel_identifikasi = [
      ("Penulis Modul", nama_penulis),
      ("Satuan Pendidikan", nama_sekolah),
      ("Mata Pelajaran", mata_pelajaran),
      ("Fase / Kelas", fase_kelas),
      (
          "Semester / Tahun Pelajaran",
          f"{semester} / {tahun_pelajaran}",
      ),
      ("Materi / Topik", topik),
      ("Alokasi Waktu", alokasi_waktu),
      ("Pertemuan Ke-", pertemuan_ke),
  ]
  add_section_table("IDENTIFIKASI DAN INFORMASI UMUM", tabel_identifikasi)

  # 2. Dimensi Profil Lulusan
  tabel_dpl = [
      (
          "Dimensi Profil Lulusan",
          data_ai.get(
              "dimensi_profil_lulusan",
              "☐ Keimanan dan Ketaqwaan terhadap Tuhan Yang Maha Esa\n"
              "☐ Kewargaan\n"
              "☑ Penalaran Kritis\n"
              "☐ Kreativitas\n"
              "☑ Kolaborasi\n"
              "☐ Kemandirian\n"
              "☐ Kesehatan\n"
              "☑ Komunikasi",
          ),
      ),
  ]
  add_section_table("DIMENSI PROFIL LULUSAN", tabel_dpl)

  # 3. Tujuan Pembelajaran
  tabel_tujuan = [
      (
          "Tujuan Pembelajaran",
          data_ai.get(
              "tujuan_pembelajaran",
              "Peserta didik mampu menguasai kompetensi sesuai materi.",
          ),
      ),
  ]
  add_section_table("TUJUAN PEMBELAJARAN", tabel_tujuan)

  # 4. Pemahaman Bermakna & Pertanyaan Pemantik
  tabel_pemahaman = [
      (
          "Pemahaman Bermakna",
          data_ai.get(
              "pemahaman_bermakna",
              "Manfaat praktis dan esensi pembelajaran bagi kehidupan.",
          ),
      ),
      (
          "Pertanyaan Pemantik",
          data_ai.get(
              "pertanyaan_pemantik",
              "Pertanyaan kritis untuk menstimulasi rasa ingin tahu peserta"
              " didik.",
          ),
      ),
  ]
  add_section_table(
      "PEMAHAMAN BERMAKNA & PERTANYAAN PEMANTIK", tabel_pemahaman
  )

  # 5. Kerangka Pembelajaran
  tabel_kerangka = [
      (
          "Praktik Pedagogis",
          data_ai.get(
              "praktik_pedagogis",
              "Model: Problem Based Learning\nMetode: Diskusi, Tanya Jawab,"
              " Analisis Teks",
          ),
      ),
      (
          "Kemitraan Pembelajaran",
          data_ai.get(
              "kemitraan_pembelajaran",
              "Lingkungan Sekolah: Kolaborasi guru mapel produktif.\nLingkungan"
              " Luar Sekolah: Pemanfaatan data/narasumber instansi terkait.",
          ),
      ),
      (
          "Lingkungan Belajar",
          data_ai.get(
              "lingkungan_belajar",
              "Ruang Fisik: Kelas fleksibel dan kolaboratif.\nRuang Virtual:"
              " Google Drive / LMS Sekolah.\nRuang/Budaya Belajar: Kolaboratif,"
              " Berpikir Kritis, Keterbukaan.",
          ),
      ),
      (
          "Pemanfaatan Digital",
          data_ai.get(
              "pemanfaatan_digital",
              "Tahap Perencanaan: AI & Cloud Storage.\nTahap Pelaksanaan: QR"
              " Code & Audio/Video Digital.\nTahap Asesmen: Google Form /"
              " Menti.",
          ),
      ),
  ]
  add_section_table("KERANGKA PEMBELAJARAN", tabel_kerangka)

  # 6. Pengalaman Belajar (Langkah-langkah)
  tabel_pengalaman = [
      (
          "Kegiatan Pendahuluan",
          data_ai.get(
              "kegiatan_pendahuluan",
              "Orientasi, Apersepsi, Motivasi, dan Asesmen Diagnostik awal.",
          ),
      ),
      (
          "Kegiatan Inti (Memahami)",
          data_ai.get(
              "kegiatan_memahami",
              "Eksplorasi konsep dan penyajian masalah autentik.",
          ),
      ),
      (
          "Kegiatan Inti (Mengaplikasi)",
          data_ai.get(
              "kegiatan_mengaplikasi",
              "Penyelidikan kolaboratif dan penerapan konsep dalam LKPD.",
          ),
      ),
      (
          "Kegiatan Inti (Merefleksi)",
          data_ai.get(
              "kegiatan_merefleksi",
              "Presentasi kelompok, umpan balik konstruktif, dan penguatan.",
          ),
      ),
      (
          "Kegiatan Penutup",
          data_ai.get(
              "kegiatan_penutup",
              "Refleksi bersama yang menyenangkan (joyful) dan bermakna.",
          ),
      ),
  ]
  add_section_table("PENGALAMAN BELAJAR (LANGKAH-LANGKAH)", tabel_pengalaman)

  # 7. Asesmen Pembelajaran
  tabel_asesmen = [
      (
          "Asesmen Awal",
          data_ai.get(
              "asesmen_awal", "Cek kesiapan sebelum masuk topik pembelajaran."
          ),
      ),
      (
          "Asesmen Proses (Formatif)",
          data_ai.get(
              "asesmen_formatif",
              "Pemantauan partisipasi, keaktifan, dan pemahaman selama"
              " kegiatan.",
          ),
      ),
      (
          "Asesmen Akhir (Sumatif)",
          data_ai.get(
              "asesmen_sumatif",
              "Evaluasi hasil berbasis unjuk kerja atau refleksi kedalaman"
              " konsep.",
          ),
      ),
  ]
  add_section_table("ASESMEN PEMBELAJARAN", tabel_asesmen)

  # 8. Rubrik Penilaian & Pedoman Penskoran
  tabel_rubrik = [
      (
          "Rubrik Penilaian",
          data_ai.get(
              "rubrik_penilaian",
              "Matriks penilaian unjuk kerja (Perlu Bimbingan, Cukup, Baik,"
              " Sangat Baik).",
          ),
      ),
      (
          "Pedoman Penskoran",
          data_ai.get(
              "pedoman_penskoran",
              "Rumus Nilai Akhir = (Skor Perolehan / Skor Maksimal) * 100",
          ),
      ),
  ]
  add_section_table("RUBRIK PENILAIAN & PEDOMAN PENSKORAN", tabel_rubrik)

  # Tanda Tangan / Pengesahan di Bagian Bawah
  p_sign = doc.add_paragraph()
  p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  p_sign.paragraph_format.space_before = Pt(14)
  p_sign.paragraph_format.space_after = Pt(4)
  p_sign.add_run(f"{nama_kota}, {tanggal_pembuatan}\nPenyusun,\n\n\n")
  run_name = p_sign.add_run(f"{nama_penulis}")
  run_name.font.bold = True
  p_sign.add_run(f"\nNIP. {nip_penulis}")

  bio = BytesIO()
  doc.save(bio)
  bio.seek(0)
  return bio


if st.button("🚀 Buat Modul Ajar Sesuai Sistematika Baru"):
  if not api_key:
    st.error("Mohon masukkan Google Gemini API Key terlebih dahulu.")
  elif not topik:
    st.warning("Mohon isi topik pembelajaran.")
  else:
    with st.spinner("Gemini sedang menyusun Modul Ajar Komprehensif..."):
      genai.configure(api_key=api_key)
      model = genai.GenerativeModel("gemini-1.5-flash")

      prompt = f"""
            Bertindaklah sebagai pakar kurikulum profesional. Buatkan konten Modul Ajar Berbasis Pembelajaran Mendalam (Deep Learning) yang **SANGAT LENGKAP, DETAIL, DAN KOMPREHENSIF** untuk:
            - Jenjang: {jenjang_pendidikan} ({fase_kelas})
            - Mata Pelajaran: {mata_pelajaran}
            - Topik / Materi Pokok: {topik}
            - Alokasi Waktu: {alokasi_waktu}
            - Pertemuan Ke-: {pertemuan_ke}

            Ketentuan Penting:
            1. Dimensi Profil Lulusan (Pilih 2 atau 4 yang paling relevan dari 8 dimensi berikut dan beri tanda centang ☑ pada yang dipilih dan ☐ pada yang tidak dipilih: Keimanan dan Ketaqwaan terhadap Tuhan Yang Maha Esa, Kewargaan, Penalaran Kritis, Kreativitas, Kolaborasi, Kemandirian, Kesehatan, Komunikasi).
            2. Praktik Pedagogis (Gunakan salah satu model: Problem Based Learning / Discovery Learning / Inquiri / Project Based Learning, serta metode pembelajaran pendukung minimal 2-3 metode).
            3. Kemitraan Pembelajaran (Lingkungan Sekolah & Lingkungan Luar Sekolah).
            4. Lingkungan Belajar (Ruang Fisik, Ruang Virtual, Ruang/Budaya Belajar).
            5. Pemanfaatan Digital (Tahap Perencanaan, Tahap Pelaksanaan, Tahap Asesmen).
            6. Pengalaman Belajar harus terstruktur mencakup Kegiatan Pendahuluan, Kegiatan Inti (Memahami, Mengaplikasi, Merefleksi), dan Kegiatan Penutup (refleksi joyful dan bermakna).
            7. Asesmen Pembelajaran mencakup Asesmen Awal, Asesmen Proses (Formatif), dan Asesmen Akhir (Sumatif) beserta Rubrik Penilaian dan Pedoman Penskorannya.

            Berikan output HANYA dalam format JSON valid yang memuat kunci-kunci berikut:
            {{
              "dimensi_profil_lulusan": "Daftar 8 DPL dengan tanda centang ☑ pada yang dipilih dan ☐ pada yang lain beserta uraian singkat penerapannya.",
              "tujuan_pembelajaran": "Uraian tujuan pembelajaran yang spesifik, operasional, dan terukur sesuai materi.",
              "pemahaman_bermakna": "Uraian pemahaman bermakna yang mendalam terkait materi.",
              "pertanyaan_pemantik": "2 pertanyaan pemantik yang kontekstual dan menantang daya nalar kritis siswa.",
              "praktik_pedagogis": "Uraian model pembelajaran (Pilih satu: Problem Based Learning / Discovery Learning / Inquiri / Project Based Learning) dan metode pembelajarannya.",
              "kemitraan_pembelajaran": "Uraian kemitraan lingkungan sekolah dan lingkungan luar sekolah secara kontekstual.",
              "lingkungan_belajar": "Uraian ruang fisik, ruang virtual, dan budaya belajar yang ingin dikembangkan.",
              "pemanfaatan_digital": "Uraian pemanfaatan digital pada tahap perencanaan, pelaksanaan, dan asesmen.",
              "kegiatan_pendahuluan": "Langkah rinci kegiatan pendahuluan (orientasi, apersepsi, asesmen awal).",
              "kegiatan_memahami": "Langkah rinci kegiatan inti pada tahap Memahami.",
              "kegiatan_mengaplikasi": "Langkah rinci kegiatan inti pada tahap Mengaplikasi.",
              "kegiatan_merefleksi": "Langkah rinci kegiatan inti pada tahap Merefleksi dan presentasi.",
              "kegiatan_penutup": "Langkah rinci kegiatan penutup yang joyful dan bermakna.",
              "asesmen_awal": "Uraian asesmen awal untuk cek kesiapan belajar.",
              "asesmen_formatif": "Uraian asesmen proses/formatif pemantauan partisipasi.",
              "asesmen_sumatif": "Uraian asesmen akhir/sumatif evaluasi unjuk kerja.",
              "rubrik_penilaian": "Matriks rubrik penilaian lengkap dari Perlu Bimbingan hingga Sangat Baik.",
              "pedoman_penskoran": "Rumus perhitungan nilai akhir dan interval predikat."
            }}
            """

      response = model.generate_content(prompt)
      text_resp = response.text.strip()

      if text_resp.startswith("```json"):
        text_resp = text_resp[7:]
      if text_resp.startswith("```"):
        text_resp = text_resp[3:]
      if text_resp.endswith("```"):
        text_resp = text_resp[:-3]
      text_resp = text_resp.strip()

      try:
        data_ai = json.loads(text_resp)
      except Exception:
        data_ai = {}

      st.success("🎉 Modul Ajar Sesuai Sistematika Baru Berhasil Disusun!")
      st.info(
          "Dokumen Word (.docx) siap diunduh dengan struktur matriks lengkap,"
          " kolom kiri berwarna, dan format rapi."
      )

      docx_file = generate_docx(
          data_ai,
          nama_sekolah,
          semester,
          tahun_pelajaran,
          mata_pelajaran,
          fase_kelas,
          topik,
          alokasi_waktu,
          pertemuan_ke,
          nama_penulis,
          nama_kota,
          tanggal_pembuatan,
          nip_penulis,
      )

      st.download_button(
          label="📥 Unduh Modul Ajar Format Tabel Matriks (.docx)",
          data=docx_file,
          file_name=f"Modul_Ajar_{topik.replace(' ', '_')}.docx",
          mime=(
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          ),
      )